# generate_stt_reports.py
# usage: python generate_stt_reports.py "/path/to/STT Manual Eval Template (2).xlsx"

import os, sys, json, math, zipfile, argparse
from datetime import datetime
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

TARGET_TABS = ["Lina","MRK","Dipto","Mehadi","Mashruf-2","Nusrat","Annoor","Annoor-2","Lina-2","Mashruf"]

def norm(s): return str(s).strip().lower().replace(" ", "_")
def safe_name(s):
    return "".join([c for c in str(s) if c.isalnum() or c in ("-","_","."," ") ]).strip().replace(" ","_")

def pick_categorical_columns(df):
    cats = []
    for c in df.columns:
        cn = norm(c)
        if df[c].dtype == "object" or df[c].dtype.name.startswith("category") or df[c].nunique(dropna=True) <= max(25, int(0.05*len(df))):
            if any(key in cn for key in ["error","issue","flag","noise","lang","domain","accent","type","source","split","pass","fail","quality","decision","status","label","comment"]):
                cats.append(c)
    if not cats:
        candidates = sorted(df.columns, key=lambda c: df[c].nunique(dropna=True))
        for c in candidates:
            if df[c].nunique(dropna=True) <= 50:
                cats.append(c)
                if len(cats) >= 3:
                    break
    # don’t include long free-text columns like “Text”
    cats = [c for c in cats if df[c].astype(str).str.len().median() < 80]
    return list(dict.fromkeys(cats))

def pick_numeric_columns(df):
    return [c for c in df.columns if pd.api.types.is_numeric_dtype(df[c])]

def plot_value_counts(series, title, out_path, topn=10):
    vc = series.astype(str).replace({"nan":"<NA>"}).value_counts().head(topn)
    if vc.empty:
        return False
    plt.figure()
    vc.plot(kind="bar")
    plt.title(title)
    plt.ylabel("Count")
    plt.tight_layout()
    plt.savefig(out_path, dpi=200)
    plt.close()
    return True

def plot_hist(series, title, out_path, bins=20):
    clean = pd.to_numeric(series, errors="coerce").dropna()
    if len(clean)==0: return False
    plt.figure()
    plt.hist(clean, bins=bins)
    plt.title(title)
    plt.xlabel(series.name)
    plt.ylabel("Frequency")
    plt.tight_layout()
    plt.savefig(out_path, dpi=200)
    plt.close()
    return True

def generate_insights(df):
    insights = []
    n = len(df)
    insights.append(f"Total items evaluated: {n:,}.")
    # Missingness
    miss = df.isna().mean().sort_values(ascending=False)
    high = miss[miss > 0.20]
    if not high.empty:
        insights.append("High missingness detected in: " + ", ".join([f"{c} ({p:.0%})" for c,p in high.head(5).items()]) + ".")
    # Dominant categories
    for c in pick_categorical_columns(df)[:4]:
        vc = df[c].astype(str).replace({"nan":"<NA>"}).value_counts(normalize=True)
        if not vc.empty:
            top_val, top_share = vc.index[0], vc.iloc[0]
            if top_share >= 0.60 and top_val != "<NA>":
                insights.append(f"'{c}' is dominated by '{top_val}' ({top_share:.0%}); potential class imbalance.")
    # Outliers in first numeric column (if any)
    for c in pick_numeric_columns(df):
        s = pd.to_numeric(df[c], errors="coerce")
        if s.notna().sum()>0:
            q1,q3 = s.quantile(0.25), s.quantile(0.75)
            iqr = q3-q1
            if iqr>0:
                out_hi = (s > q3 + 1.5*iqr).mean()
                if out_hi>0.05:
                    insights.append(f"'{c}' shows {out_hi:.0%} high-end outliers; review or cap.")
            break
    return insights

def generate_recommendations(df):
    recs = []
    cols = [norm(c) for c in df.columns]
    if any("noise" in c for c in cols):
        recs.append("Filter/down-weight high-noise segments using the noise flags; consider noise profiling.")
    else:
        recs.append("Introduce a standardized noise_level flag (0–3) to guide filtering and augmentation.")
    if any("overlap" in c or "cross_talk" in c or "crosstalk" in c for c in cols):
        recs.append("Handle cross-talk via diarization or exclusion; avoid heavy overlap in training batches.")
    if any("lang" in c for c in cols):
        recs.append("Tag code-mixed content explicitly; exclude or model it separately if needed.")
    if any("duration" in c or "length" in c or "secs" in c for c in cols):
        recs.append("Normalize clip durations (e.g., 8–20s) or bucket them for stable batching and convergence.")
    recs.append("Standardize transcript normalization (digits, punctuation, casing, diacritics) with one shared pipeline.")
    recs.append("Run labeler calibration rounds with rubric; track pass/fail and issue taxonomy per annotator weekly.")
    return recs[:6]

def build_path_map(count_df):
    path_map = {}
    if count_df.empty:
        return path_map
    tab_cols = [c for c in count_df.columns if any(k in c for k in ["tab","sheet","name","annotator"])]
    path_cols = [c for c in count_df.columns if any(k in c for k in ["path","folder","directory"])]
    if not tab_cols or not path_cols:
        return path_map
    tcol, pcol = tab_cols[0], path_cols[0]
    for _, row in count_df.iterrows():
        tval = str(row.get(tcol, "")).strip()
        pval = str(row.get(pcol, "")).strip()
        if tval:
            path_map[tval] = pval
    return path_map

def write_docx(tab_name, path_for_tab, df, image_paths, out_dir):
    doc = Document()
    title = f"Evaluation Report by ({tab_name})"
    subtitle = f"Evaluation of data in {path_for_tab}" if path_for_tab else "Evaluation of data (path unknown)"
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(subtitle); p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Overview
    doc.add_heading("1. Overview", level=1)
    doc.add_paragraph(
        "This report summarizes manual evaluations of Bangla Speech-to-Text training data. "
        "Each item corresponds to an audio file stored under a 'flac' folder with a paired transcript in a JSON file. "
        "We provide dataset-level summaries, charts, and actionable recommendations."
    )

    # Snapshot
    doc.add_heading("2. Dataset Snapshot", level=1)
    n = len(df)
    doc.add_paragraph(f"Total evaluated items: {n:,}")
    doc.add_paragraph("Columns detected: " + ", ".join(map(str, df.columns.tolist())))

    # Categorical
    cat_cols = pick_categorical_columns(df)
    if cat_cols:
        doc.add_heading("3. Key Categorical Distributions", level=1)
        for c in cat_cols[:4]:
            vc = df[c].astype(str).replace({"nan":"<NA>"}).value_counts().head(10)
            doc.add_paragraph(f"Top values for '{c}':")
            t = doc.add_table(rows=1, cols=2)
            t.rows[0].cells[0].text = "Value"
            t.rows[0].cells[1].text = "Count"
            for val, cnt in vc.items():
                row = t.add_row().cells
                row[0].text, row[1].text = str(val), str(cnt)
            doc.add_paragraph("")

    # Numeric
    num_cols = pick_numeric_columns(df)
    if num_cols:
        doc.add_heading("4. Numeric Fields Summary", level=1)
        desc = df[num_cols].describe().T
        t = doc.add_table(rows=1, cols=desc.shape[1]+1)
        header = ["Field"] + desc.columns.tolist()
        for j, col in enumerate(header):
            t.rows[0].cells[j].text = col
        for idx in desc.index:
            row = t.add_row().cells
            row[0].text = str(idx)
            for j, col in enumerate(desc.columns, start=1):
                val = desc.loc[idx, col]
                row[j].text = f"{val:.3f}" if isinstance(val, (int,float,np.floating)) else str(val)

    # Charts
    if image_paths:
        doc.add_heading("5. Charts", level=1)
        for img in image_paths:
            doc.add_paragraph(os.path.basename(img))
            doc.add_picture(img, width=Inches(5.8))

    # Insights + Recs
    doc.add_heading("6. Insights", level=1)
    for ins in generate_insights(df):
        doc.add_paragraph(f"• {ins}")
    doc.add_heading("7. Recommendations", level=1)
    for rec in generate_recommendations(df):
        doc.add_paragraph(f"• {rec}")

    path_component = safe_name(path_for_tab) if path_for_tab else safe_name(tab_name)
    doc_name = f"{safe_name(tab_name)}__{path_component}__STT_Eval_Report.docx"
    doc_path = os.path.join(out_dir, doc_name)
    doc.save(doc_path)
    return doc_path

def run_analysis(xlsx_path):
    xl = pd.ExcelFile(xlsx_path)
    # Count tab for mapping
    count_name = None
    for nm in xl.sheet_names:
        if nm.strip().lower()=="count":
            count_name = nm; break
    count_df = pd.read_excel(xlsx_path, sheet_name=count_name) if count_name else pd.DataFrame()
    # normalize columns for mapping
    if not count_df.empty:
        count_df.columns = [norm(c) for c in count_df.columns]
    path_map = build_path_map(count_df)
    path_map.setdefault("Lina", "collect/miscellaneous/Dhaka FM")  # per your rule

    out_root = "STT_eval_reports_" + datetime.now().strftime("%Y%m%d_%H%M%S")
    os.makedirs(out_root, exist_ok=True)

    generated = []
    for tab in TARGET_TABS:
        # tolerant match
        tab_effective = tab if tab in xl.sheet_names else next((s for s in xl.sheet_names if s.strip().lower()==tab.strip().lower()), None)
        if tab_effective is None:
            continue
        df = pd.read_excel(xlsx_path, sheet_name=tab_effective)

        tab_dir = os.path.join(out_root, safe_name(tab))
        os.makedirs(tab_dir, exist_ok=True)

        # plots
        images = []
        for c in pick_categorical_columns(df)[:3]:
            out_img = os.path.join(tab_dir, f"{safe_name(tab)}__{safe_name(c)}__value_counts.png")
            if plot_value_counts(df[c], f"{tab}: {c} distribution (Top 10)", out_img):
                images.append(out_img)
        hist_count = 0
        for c in pick_numeric_columns(df):
            out_img = os.path.join(tab_dir, f"{safe_name(tab)}__{safe_name(c)}__hist.png")
            if plot_hist(df[c], f"{tab}: {c} histogram", out_img, bins=20):
                images.append(out_img); hist_count += 1
                if hist_count>=2: break

        path_for_tab = path_map.get(tab, path_map.get(tab_effective, ""))

        docx_path = write_docx(tab, path_for_tab, df, images, tab_dir)

        # Per-tab ZIP (avoid master ZIP size/time issues)
        zip_path = os.path.join(out_root, f"{safe_name(tab)}.zip")
        with zipfile.ZipFile(zip_path, 'w', compression=zipfile.ZIP_DEFLATED, allowZip64=True) as zf:
            zf.write(docx_path, arcname=os.path.join(safe_name(tab), os.path.basename(docx_path)))
            for img in images:
                zf.write(img, arcname=os.path.join(safe_name(tab), "images", os.path.basename(img)))

        generated.append({
            "tab": tab,
            "docx": os.path.abspath(docx_path),
            "zip": os.path.abspath(zip_path),
            "images": [os.path.abspath(p) for p in images],
            "path_from_count": path_for_tab
        })

    with open(os.path.join(out_root, "summary.json"), "w", encoding="utf-8") as f:
        json.dump(generated, f, indent=2, ensure_ascii=False)

    print(f"Done. Output folder: {os.path.abspath(out_root)}")
    for g in generated:
        print(f"- {g['tab']}: {g['zip']}")

if __name__ == "__main__":
    #ap = argparse.ArgumentParser()
    #ap.add_argument("xlsx_path", help="Path to the Excel file (e.g., STT Manual Eval Template (2).xlsx)")
    #args = ap.parse_args()
    xlsx_path = './STT Manual Eval Template (1).xlsx'
    run_analysis(xlsx_path)
    #main(args.xlsx_path)
